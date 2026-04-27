"""文档问答（PDF/Word）。

演示 LangChain 概念：
- Document Loaders：PyPDFLoader、DocxLoader
- 多格式文档统一处理
- 文档切分 + 向量检索 + 问答

对照学习：
- 现有 RAG：只支持 .md/.txt 纯文本
- 文档问答：支持 PDF/Word 等真实文档格式
- Document Loader 是 LangChain 生态的重要扩展点
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.tools import StructuredTool
from langchain_core.vectorstores import InMemoryVectorStore
from pydantic import BaseModel, Field

from ..config import Settings
from ..logging_utils import get_logger
from ..openai_support import build_chat_model, build_embeddings, ensure_chat_api_key, ensure_embedding_api_key

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class DocumentQuestionInput(BaseModel):
    file_path: str = Field(description="文档文件路径")
    question: str = Field(description="要问的问题")


def load_document(file_path: str) -> list[Document]:
    """根据文件扩展名选择对应的 Loader 加载文档。"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = path.suffix.lower()
    logger.info("加载文档: path=%s, ext=%s", file_path, ext)

    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(Document(
                    page_content=text,
                    metadata={"source": str(path), "page": i + 1},
                ))
        logger.info("PDF 加载完成: pages=%s", len(pages))
        return pages

    elif ext == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append(Document(
                    page_content=para.text,
                    metadata={"source": str(path), "paragraph": i + 1},
                ))
        logger.info("Word 文档加载完成: paragraphs=%s", len(paragraphs))
        return paragraphs

    elif ext in (".txt", ".md"):
        text = path.read_text(encoding="utf-8")
        logger.info("文本文件加载完成: chars=%s", len(text))
        return [Document(page_content=text, metadata={"source": str(path)})]

    else:
        raise ValueError(f"不支持的文件格式: {ext}（支持: {', '.join(SUPPORTED_EXTENSIONS)}）")


def split_into_chunks(documents: list[Document], chunk_size: int = 800, chunk_overlap: int = 120) -> list[Document]:
    """将文档切分为固定大小的块。"""
    chunks = []
    for doc in documents:
        text = doc.page_content
        if len(text) <= chunk_size:
            chunks.append(doc)
            continue
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            chunks.append(Document(
                page_content=chunk_text,
                metadata={**doc.metadata, "chunk": len(chunks)},
            ))
            start += chunk_size - chunk_overlap
    logger.info("文档切分完成: chunks=%s", len(chunks))
    return chunks


def answer_document_question(
    file_path: str,
    question: str,
    settings: Settings,
    *,
    top_k: int = 4,
) -> str:
    """加载文档并回答问题。"""
    ensure_chat_api_key(settings)
    ensure_embedding_api_key(settings)

    documents = load_document(file_path)
    if not documents:
        return "文档内容为空，无法回答问题。"

    chunks = split_into_chunks(documents)
    embeddings = build_embeddings(settings)
    vector_store = InMemoryVectorStore(embeddings)
    vector_store.add_documents(chunks)

    retrieved = vector_store.similarity_search(question, k=min(top_k, len(chunks)))
    logger.info("文档检索完成: retrieved=%s", len(retrieved))

    context = "\n\n".join(
        f"[来源: {doc.metadata.get('source', '?')}, 页码/段落: {doc.metadata.get('page', doc.metadata.get('paragraph', '?'))}]\n{doc.page_content}"
        for doc in retrieved
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", "你是一个文档问答助手。只根据提供的文档内容回答，不要编造。若文档不足，请说明。用中文回答。"),
        ("human", "文档内容：\n{context}\n\n问题：{question}"),
    ])

    chain = prompt | build_chat_model(settings) | StrOutputParser()
    return chain.invoke({"context": context, "question": question})


def build_document_qa_tool(settings: Settings) -> StructuredTool:
    """构建文档问答工具（供 Agent 使用）。"""

    def doc_qa(file_path: str, question: str) -> str:
        logger.info("document_qa 工具被调用: file=%s, question=%s", file_path, question[:80])
        try:
            return answer_document_question(file_path, question, settings)
        except Exception as exc:
            return f"文档处理失败: {exc}"

    return StructuredTool.from_function(
        func=doc_qa,
        name="document_qa",
        description="读取 PDF/Word/TXT 文件并回答问题。需要文件路径和问题。",
        args_schema=DocumentQuestionInput,
    )
